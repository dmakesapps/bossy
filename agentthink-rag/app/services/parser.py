"""
Document Parser Service

Handles parsing of various document formats: PDF, DOCX, CSV, TXT, MD.
Extracts text content with metadata for downstream processing.
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Any

from app.utils.text_cleaner import clean_text, is_meaningful_text

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    """Represents a section of parsed content from a document."""
    
    content: str
    """The extracted text content."""
    
    page_number: Optional[int] = None
    """Page number for PDFs, None for other formats."""
    
    section_title: Optional[str] = None
    """Section heading for structured documents."""
    
    doc_type: str = "unknown"
    """Document type: pdf, docx, csv, txt, md."""
    
    metadata: dict = field(default_factory=dict)
    """Additional metadata: heading_level, table data, etc."""


class DocumentParser:
    """
    Multi-format document parser.
    
    Supports PDF, DOCX, CSV, TXT, and Markdown files.
    Extracts text with metadata for each section.
    """
    
    def __init__(self):
        """Initialize the parser."""
        self._parsers = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'csv': self._parse_csv,
            'txt': self._parse_txt,
            'md': self._parse_markdown,
        }
    
    def parse(self, file_path: str, file_type: str) -> list[ParsedSection]:
        """
        Parse a document and extract sections.
        
        Args:
            file_path: Path to the document file
            file_type: File extension (pdf, docx, csv, txt, md)
            
        Returns:
            List of ParsedSection objects containing extracted content
        """
        file_type = file_type.lower().lstrip('.')
        
        if file_type not in self._parsers:
            logger.error(f"Unsupported file type: {file_type}")
            return []
        
        filepath = Path(file_path)
        if not filepath.exists():
            logger.error(f"File not found: {file_path}")
            return []
        
        try:
            parser = self._parsers[file_type]
            sections = parser(file_path)
            
            # Filter out empty/meaningless sections
            sections = [s for s in sections if is_meaningful_text(s.content)]
            
            logger.info(f"Parsed {filepath.name}: {len(sections)} sections extracted")
            return sections
            
        except Exception as e:
            logger.exception(f"Error parsing {file_path}: {e}")
            return []
    
    def _parse_pdf(self, file_path: str) -> list[ParsedSection]:
        """
        Parse a PDF file using PyMuPDF with pdfplumber fallback for tables.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of ParsedSection objects
        """
        import fitz  # PyMuPDF
        
        sections = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc, start=1):
                # Primary extraction using get_text("text")
                text = page.get_text("text")
                
                # If text is messy or appears limited, use blocks extraction with sorting
                # for more robust reading order (crucial for unformatted/fragmented PDFs)
                blocks = page.get_text("blocks")
                # Sort blocks: vertically (y1 coordinate is blocks[1]), then horizontally (x0 is blocks[0])
                blocks.sort(key=lambda b: (b[1], b[0]))
                
                sorted_text_parts = []
                for b in blocks:
                    if len(b) >= 5 and isinstance(b[4], str):
                        content = b[4].strip()
                        if content:
                            sorted_text_parts.append(content)
                
                sorted_text = "\n".join(sorted_text_parts)
                
                # Use the sorted text if it's more substantial or as primary
                if len(sorted_text) > len(text) * 0.9:
                    text = sorted_text
                    
                # Clean the extracted text
                text = clean_text(text)
                
                if text and is_meaningful_text(text):
                    sections.append(ParsedSection(
                        content=text,
                        page_number=page_num,
                        doc_type="pdf",
                        metadata={"extraction_method": "pymupdf"}
                    ))
                else:
                    logger.warning(f"Page {page_num} appears to be scanned/image-only")
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PyMuPDF error on {file_path}: {e}")
        
        # Try to extract tables using pdfplumber
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    
                    for table_idx, table in enumerate(tables):
                        if not table or not table[0]:
                            continue
                        
                        # Convert table to text format
                        table_text = self._table_to_text(table, page_num, table_idx)
                        
                        if table_text and is_meaningful_text(table_text):
                            sections.append(ParsedSection(
                                content=table_text,
                                page_number=page_num,
                                doc_type="pdf",
                                metadata={
                                    "type": "table",
                                    "table_index": table_idx,
                                    "extraction_method": "pdfplumber"
                                }
                            ))
                            
        except Exception as e:
            logger.warning(f"pdfplumber table extraction failed: {e}")
        
        return sections
    
    def _parse_docx(self, file_path: str) -> list[ParsedSection]:
        """
        Parse a DOCX file using python-docx.
        
        Groups paragraphs under headings into sections.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of ParsedSection objects
        """
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        
        sections = []
        doc = Document(file_path)
        
        current_heading = None
        current_heading_level = None
        current_paragraphs = []
        
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            
            # Check if this is a heading
            if style_name.startswith("Heading"):
                # Save previous section if exists
                if current_paragraphs:
                    content = "\n\n".join(current_paragraphs)
                    content = clean_text(content)
                    
                    if is_meaningful_text(content):
                        sections.append(ParsedSection(
                            content=content,
                            section_title=current_heading,
                            doc_type="docx",
                            metadata={"heading_level": current_heading_level}
                        ))
                
                # Start new section
                current_heading = para.text.strip()
                try:
                    current_heading_level = int(style_name.replace("Heading ", "").split()[0])
                except (ValueError, IndexError):
                    current_heading_level = 1
                current_paragraphs = []
            else:
                # Regular paragraph
                text = para.text.strip()
                if text:
                    current_paragraphs.append(text)
        
        # Don't forget the last section
        if current_paragraphs:
            content = "\n\n".join(current_paragraphs)
            content = clean_text(content)
            
            if is_meaningful_text(content):
                sections.append(ParsedSection(
                    content=content,
                    section_title=current_heading,
                    doc_type="docx",
                    metadata={"heading_level": current_heading_level}
                ))
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                table_text = self._table_list_to_text(table_data, table_idx)
                
                if is_meaningful_text(table_text):
                    sections.append(ParsedSection(
                        content=table_text,
                        doc_type="docx",
                        metadata={"type": "table", "table_index": table_idx}
                    ))
        
        return sections
    
    def _parse_csv(self, file_path: str) -> list[ParsedSection]:
        """
        Parse a CSV file using pandas.
        
        Converts to text in groups of 20 rows per section.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            List of ParsedSection objects
        """
        import pandas as pd
        
        sections = []
        
        try:
            # Try UTF-8 first, then fallback to latin-1
            try:
                df = pd.read_csv(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, encoding='latin-1', errors='replace')
            
            total_rows = len(df)
            columns = df.columns.tolist()
            columns_str = ", ".join(str(col) for col in columns)
            
            # Process in groups of 20 rows
            batch_size = 20
            
            for start_idx in range(0, total_rows, batch_size):
                end_idx = min(start_idx + batch_size, total_rows)
                batch = df.iloc[start_idx:end_idx]
                
                # Build text representation
                text_parts = [f"Columns: {columns_str}"]
                
                for row_num, (_, row) in enumerate(batch.iterrows(), start=start_idx + 1):
                    row_values = ", ".join(str(v) for v in row.values)
                    text_parts.append(f"Row {row_num}: {row_values}")
                
                content = "\n".join(text_parts)
                content = clean_text(content)
                
                if is_meaningful_text(content):
                    sections.append(ParsedSection(
                        content=content,
                        doc_type="csv",
                        metadata={
                            "row_range": f"{start_idx + 1}-{end_idx}",
                            "total_rows": total_rows,
                            "columns": columns
                        }
                    ))
                    
        except Exception as e:
            logger.error(f"Error parsing CSV {file_path}: {e}")
        
        return sections
    
    def _parse_txt(self, file_path: str) -> list[ParsedSection]:
        """
        Parse a plain text file.
        
        Splits into sections by paragraph breaks (double newlines).
        
        Args:
            file_path: Path to the text file
            
        Returns:
            List of ParsedSection objects
        """
        sections = []
        content = None
        
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            logger.error(f"Could not decode text file: {file_path}")
            return []
        
        # Split by paragraph breaks
        paragraphs = content.split('\n\n')
        
        for idx, para in enumerate(paragraphs):
            para = clean_text(para)
            
            if is_meaningful_text(para):
                sections.append(ParsedSection(
                    content=para,
                    doc_type="txt",
                    metadata={"paragraph_index": idx}
                ))
        
        return sections
    
    def _parse_markdown(self, file_path: str) -> list[ParsedSection]:
        """
        Parse a Markdown file.
        
        Splits at heading boundaries (lines starting with #).
        
        Args:
            file_path: Path to the Markdown file
            
        Returns:
            List of ParsedSection objects
        """
        import re
        
        sections = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
        
        # Split at heading boundaries
        # Pattern matches lines starting with one or more #
        heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        # Find all headings with their positions
        matches = list(heading_pattern.finditer(content))
        
        if not matches:
            # No headings, treat entire content as one section
            cleaned = clean_text(content)
            if is_meaningful_text(cleaned):
                sections.append(ParsedSection(
                    content=cleaned,
                    doc_type="md",
                    metadata={}
                ))
            return sections
        
        # Process content before first heading
        if matches[0].start() > 0:
            preamble = content[:matches[0].start()]
            preamble = clean_text(preamble)
            if is_meaningful_text(preamble):
                sections.append(ParsedSection(
                    content=preamble,
                    doc_type="md",
                    metadata={"is_preamble": True}
                ))
        
        # Process each heading section
        for i, match in enumerate(matches):
            heading_level = len(match.group(1))
            heading_text = match.group(2).strip()
            
            # Get content until next heading or end
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            
            section_content = content[start:end]
            section_content = clean_text(section_content)
            
            if is_meaningful_text(section_content):
                sections.append(ParsedSection(
                    content=section_content,
                    section_title=heading_text,
                    doc_type="md",
                    metadata={"heading_level": heading_level}
                ))
        
        return sections
    
    def _table_to_text(self, table: list, page_num: int, table_idx: int) -> str:
        """Convert a pdfplumber table to text format."""
        if not table:
            return ""
        
        lines = [f"Table on Page {page_num}:"]
        
        for row in table:
            if row:
                # Clean each cell value
                row_cells = [str(cell).strip() if cell else "" for cell in row]
                lines.append("| " + " | ".join(row_cells) + " |")
        
        return "\n".join(lines)
    
    def _table_list_to_text(self, table_data: list[list[str]], table_idx: int) -> str:
        """Convert a list of rows to text format."""
        if not table_data:
            return ""
        
        lines = [f"Table {table_idx + 1}:"]
        
        for row in table_data:
            if row:
                lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(lines)
